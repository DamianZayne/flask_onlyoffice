"""
智慧办案助手 — Flask 应用
支持 OnlyOffice 离线文档编辑
"""

import json
import os
import re
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path

import requests
from docx import Document
from flask import (
    Flask,
    abort,
    g,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    url_for,
)

# ── 配置 ──────────────────────────────────────────────────────
# BASE_DIR = Path(__file__).resolve().parent

import sys
if getattr(sys, 'frozen', False):
    BASE_DIR = Path(sys._MEIPASS)
else:
    BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
META_FILE = BASE_DIR / "meta.json"

MIME_MAP = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "pdf":  "application/pdf",
    "doc":  "application/msword",
    "xls":  "application/vnd.ms-excel",
    "ppt":  "application/vnd.ms-powerpoint",
    "odt":  "application/vnd.oasis.opendocument.text",
    "ods":  "application/vnd.oasis.opendocument.spreadsheet",
    "odp":  "application/vnd.oasis.opendocument.presentation",
    "csv":  "text/csv",
    "txt":  "text/plain",
}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024

# ── SQLite 数据库 ─────────────────────────────────────────────
DB_PATH = BASE_DIR / "data.db"

#Flask 应用上下文全局数据库连接管理
#
'''
g 是 Flask 的应用上下文全局变量
在同一个请求中，所有代码共享同一个 g 对象
请求结束后，g 对象会被销毁
适合存储每个请求需要复用的资源（如数据库连接）'''
def get_db():
    if "db" not in g:# 检查应用上下文全局变量中是否有数据库连接
        g.db = sqlite3.connect(str(DB_PATH))#创建数据库连接
        g.db.row_factory = sqlite3.Row# 设置行工厂，返回字典样式的行
        g.db.execute("PRAGMA journal_mode=WAL")
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db
'''
journal_mode=WAL：Write-Ahead Logging 模式

允许并发读写
提高写入性能
减少锁定
foreign_keys=ON：启用外键约束

确保数据完整性
自动级联操作

'''

@app.teardown_appcontext
def close_db(exception):
    db = g.pop("db", None)
    if db:
        db.close()


def init_db():
    db = sqlite3.connect(str(DB_PATH))
    db.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_name TEXT DEFAULT '',
            subject_name TEXT DEFAULT '',
            unit_name TEXT DEFAULT '',
            archivist TEXT DEFAULT '',
            archive_date TEXT DEFAULT '',
            status TEXT DEFAULT 'active',
            reviewer TEXT DEFAULT '',
            created_at TEXT DEFAULT ''
        )
    """)
    db.execute("""
        CREATE TABLE IF NOT EXISTS placeholders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id INTEGER NOT NULL,
            param_set TEXT NOT NULL DEFAULT 'A',
            code TEXT NOT NULL,
            value TEXT DEFAULT '',
            FOREIGN KEY (case_id) REFERENCES cases(id) ON DELETE CASCADE
        )
    """)
    db.execute('''
            CREATE TABLE IF NOT EXISTS template_mappings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_name TEXT NOT NULL,
                filename TEXT NOT NULL,
                dir_num INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(group_name, filename)  -- 同一个组内文件名唯一
            )
        ''')
    db.commit()
    db.close()



init_db()

# ── 材料名 → Edit_Doc 目录映射 ────────────────────────────────
EDIT_DOC_DIR = BASE_DIR / "Edit_Doc"


# ── 元数据 ────────────────────────────────────────────────────
def load_meta():
    if META_FILE.exists():
        return json.loads(META_FILE.read_text(encoding="utf-8"))#//整个JSON文件的所有内容
    return {}

def save_meta(meta):
    META_FILE.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


# ── SDK 静态资源 ──────────────────────────────────────────────
@app.route("/packages/<path:filename>")
def serve_packages(filename):
    return send_from_directory(BASE_DIR / "static" / "packages", filename)#安全地发送静态文件的函数，从指定目录安全地提供文件下载


# ── 业务页面 ──────────────────────────────────────────────────

@app.route("/")
def index():
    """首页 - 工作台"""
    return render_template("首页.html")


@app.route("/login")
def login():
    """登录页"""
    return render_template("登录.html")


@app.route("/security-catalog")
def security_catalog():
    """安全卷录入"""
    return render_template("安全卷录入.html")


# ── OnlyOffice 编辑器 ──────────────    ───────────────────────────

@app.route("/editor/<doc_id>")
def edit_doc(doc_id):
    """编辑已上传的文档"""
    meta = load_meta()
    info = meta.get(doc_id)
    if not info:
        abort(404)

    filepath = UPLOAD_DIR / info["filename"]
    if not filepath.exists():
        abort(404)

    return render_template(
        "editor.html",
        doc_id=doc_id,
        doc_name=info["name"],
        doc_ext=info["ext"],
        doc_url=url_for("download_doc", doc_id=doc_id, _external=True),
        save_url=url_for("save_doc", doc_id=doc_id, _external=True),
        back_url="/",
    )


@app.route("/editor/material")
def edit_material():
    """编辑安全卷材料文档 —— 优先加载 Edit_Doc 里预置的模板"""
    material_name = request.args.get("name", "未命名材料")#获取名为 "name" 的参数值
#如果参数不存在，返回默认值 "未命名材料
    case_id = request.args.get("id", "new")
    group=request.args.get("group", "")
    back_url=f"/security-catalog?id={case_id}&group={group}"

    # 查找 Edit_Doc 中对应的文档（先查内置映射，再查用户模板）
    dir_num = ""
    print(dir_num)
    if not dir_num:
        # 用户模板：文件名（去掉 .docx）匹配
        all_t = get_all_templates(group)
        print(all_t)
        for fn, dn in all_t.items():
            display = get_material_display_name(fn)

            if display == material_name or fn.replace(".docx", "") == material_name:
                dir_num = dn
                break
    doc_url = ""
    print(dir_num)
    doc_ext = "docx"
    doc_name = f"{material_name}.docx"

    # 案件专属目录: Edit_Doc/{case_id}/{dir_num}/
    if dir_num:
        case_doc_dir = EDIT_DOC_DIR / str(case_id) / str(dir_num)
        if case_doc_dir.exists():
            files = sorted(case_doc_dir.iterdir(), key=lambda f: (f.suffix != ".docx", f.name))
            if files:
                doc_file = files[0]  # 优先 .docx
                doc_ext = doc_file.suffix.lstrip(".").lower()
                doc_name = doc_file.name
                doc_url = url_for("serve_edit_case_doc", case_id=case_id, dir_num=dir_num,
                                  filename=doc_file.name, _external=True)
    print("查找材料文档:", material_name, doc_url)

    save_url = ""
    if dir_num and doc_url:
        save_url = url_for("save_edit_case_doc", case_id=case_id, dir_num=dir_num,
                           filename=doc_name, _external=True)

    return render_template(
        "editor.html",
        doc_id=f"material_{case_id}_{dir_num or 'new'}",
        doc_name=doc_name,
        doc_ext=doc_ext,
        doc_url=doc_url,
        save_url=save_url,
        back_url=back_url,
    )


@app.route("/edit-doc/<int:case_id>/<int:dir_num>/<path:filename>")
def serve_edit_case_doc(case_id, dir_num, filename):
    """提供案件专属 Edit_Doc 目录下的文档"""
    filepath = EDIT_DOC_DIR / str(case_id) / str(dir_num) / filename
    if not filepath.exists():
        abort(404)
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "docx"
    return send_file(filepath, as_attachment=False,
                     mimetype=MIME_MAP.get(ext, "application/octet-stream"))


# ── 文档 API ─────────────────────────────────────────────────
#
# @app.route("/api/doc/<doc_id>/download")
# def download_doc(doc_id):
#     """下载已上传的文档"""
#     meta = load_meta()
#     info = meta.get(doc_id)
#     if not info:
#         abort(404)
#
#     filepath = UPLOAD_DIR / info["filename"]
#     if not filepath.exists():
#         abort(404)
#
#     return send_file(filepath, as_attachment=False, mimetype=MIME_MAP.get(info["ext"], "application/octet-stream"))
#
#
# @app.route("/api/doc/<doc_id>/info")
# def doc_info(doc_id):
#     meta = load_meta()
#     info = meta.get(doc_id)
#     if not info:
#         abort(404)
#     return jsonify(info)


# ── 保存接口 ─────────────────────────────────────────────────

@app.route("/api/save-doc/<doc_id>", methods=["POST"])
def save_doc(doc_id):
    """保存已上传文档的修改"""
    meta = load_meta()
    info = meta.get(doc_id)
    if not info:
        abort(404)

    filepath = UPLOAD_DIR / info["filename"]
    filepath.write_bytes(request.get_data())
    print(f"[save] 已保存: {filepath}")
    return jsonify({"error": 0})


@app.route("/api/save-edit-doc/<int:case_id>/<int:dir_num>/<path:filename>", methods=["POST"])
def save_edit_case_doc(case_id, dir_num, filename):
    """保存案件专属 Edit_Doc 材料文档的修改"""
    doc_dir = EDIT_DOC_DIR / str(case_id) / str(dir_num)
    # 优先覆写 .docx 文件
    existing = sorted(doc_dir.glob("*"), key=lambda f: (f.suffix != ".docx", f.name)) if doc_dir.exists() else []
    target = existing[0] if existing else doc_dir / filename
    target.write_bytes(request.get_data())
    print(f"[save-edit-doc] case={case_id} dir={dir_num}: {target}")
    return jsonify({"error": 0})


# ── 模板替换（全部来自用户上传）─────────────────────────────

@app.route("/api/replace-from-temp", methods=["POST"])
def replace_from_temp():
    """从 temp_docx 读取模板，替换占位符，写出到 Edit_Doc"""
    print("""从 temp_docx 读取模板，替换占位符，写出到 Edit_Doc""")


    data = request.get_json(force=True)
    replacements = data.get("replacements", {})
    active_set = data.get("set", "A")
    case_id = data.get("case_id", 0)

    if not replacements:
        return jsonify({"error": 1, "message": "没有替换数据"}), 400
    param_config  =  load_param_groups()
    all_templates = get_all_templates(active_set)
    selected_docs = param_config.get("doc_selections", {}).get(active_set, list(all_templates.keys()))
    print("参数组:", active_set, "已选文档:", selected_docs)

    # 只保留有值的替换项
    filled = {k: v for k, v in replacements.items() if v}
    if not filled:
        return jsonify({"error": 1, "message": "所有占位符都为空"}), 400

    results = []
    for filename, dir_num in all_templates.items():
        print("正在处理:", filename)
        if filename not in selected_docs:
            results.append({"file": filename, "status": "skip", "reason": "该参数组未选中此文档"})
            continue

        src = USER_TEMPLATES_DIR / active_set/filename
        print("源文件:", src)

        if not src.exists():
            results.append({"file": filename, "status": "skip", "reason": "源文件不存在"})
            continue
        if src.suffix.lower() != ".docx":
            results.append({"file": filename, "status": "skip", "reason": "非docx格式"})
            continue

        try:
            print("正在处理src:", src)
            doc = Document(str(src))
            # print("正在处理src:", src)

            # 扫描文档中实际使用的占位符
            import re as _re
            doc_placeholders = set()
            for p in doc.paragraphs:
                doc_placeholders.update(_re.findall(r'\{\{[^}]+\}\}', p.text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            doc_placeholders.update(_re.findall(r'\{\{[^}]+\}\}', p.text))

            # 检查 filled 中是否有该文档需要的占位符
            has_placeholders = bool(doc_placeholders)
            if not has_placeholders or not (doc_placeholders & set(filled.keys())):
                # 无占位符或不需要替换：直接复制到 Edit_Doc，不做替换
                case_id = data.get("case_id", 0)
                dst_dir = EDIT_DOC_DIR / str(case_id) / str(dir_num)
                dst_dir.mkdir(parents=True, exist_ok=True)
                dst = dst_dir / filename
                dst.write_bytes(src.read_bytes())
                reason = "无占位符，直接复制" if not has_placeholders else "未填写该文档需要的占位符，直接复制"
                results.append({"file": filename, "dir_num": dir_num, "status": "ok", "replaced": 0, "note": reason})
                print(f"[replace-from-temp] {filename} → Edit_Doc/{case_id}/{dir_num}/ (直接复制)")
                continue




            # 执行替换，保留每个 run 的格式
            import re as _re2

            def replace_in_paragraph(paragraph, reps):
                runs = paragraph.runs
                if not runs:
                    return

                # 先尝试单 run 内替换（保留格式）
                for run in runs:
                    for old_text, new_text in reps.items():
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

                # 再检查是否还有跨 run 的占位符残留
                full = paragraph.text
                remaining = _re2.findall(r'\{\{[^}]+\}\}', full)
                if not remaining:
                    return

                # 跨 run 替换：定位每个残留占位符在全文中的位置，找到对应的 run 并替换
                for placeholder in remaining:
                    if placeholder not in reps:
                        continue
                    new_val = reps[placeholder]
                    # 找到占位符在全文中的位置
                    pos = full.find(placeholder)
                    if pos < 0:
                        continue
                    end_pos = pos + len(placeholder)

                    # 找到覆盖这个位置范围的所有 run
                    run_start = 0
                    affected_runs = []
                    for run in runs:
                        run_end = run_start + len(run.text)
                        if run_end > pos and run_start < end_pos:
                            affected_runs.append((run, run_start))
                        run_start = run_end

                    if not affected_runs:
                        continue

                    # 拼接受影响 run 的文本
                    combined = "".join(r.text for r, _ in affected_runs)
                    # 在拼接文本中替换
                    rel_pos = pos - affected_runs[0][1]  # 占位符在 combined 中的起始位置
                    rel_end = rel_pos + len(placeholder)
                    combined = combined[:rel_pos] + new_val + combined[rel_end:]

                    # 写回各 run：替换值写入最后一个受影响的 run（通常带格式），其余清空
                    last_run = affected_runs[-1][0]
                    last_run.text = combined
                    for r, _ in affected_runs[:-1]:
                        r.text = ""

                    # 更新全文缓存
                    full = paragraph.text

            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, replacements)
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

            # 保存到 Edit_Doc/{case_id}/{dir_num}/
            dst_dir = EDIT_DOC_DIR / str(case_id) / str(dir_num)
            print(f"[replace-from-temp] {filename} → Edit_Doc/{case_id}/{dir_num}/")
            dst_dir.mkdir(parents=True, exist_ok=True)
            dst = dst_dir / filename
            doc.save(str(dst))

            replaced_count = len(doc_placeholders & set(filled.keys()))
            results.append({"file": filename, "dir_num": dir_num, "status": "ok",
                            "replaced": replaced_count})
            print(f"[replace-from-temp] {filename} → Edit_Doc/{case_id}/{dir_num}/, 替换 {replaced_count} 处")

        except Exception as e:
            results.append({"file": filename, "status": "error", "reason": str(e)})
            print(f"[replace-from-temp] {filename} 失败: {e}")

    ok_count = sum(1 for r in results if r["status"] == "ok")
    return jsonify({
        "error": 0,
        "message": f"成功替换 {ok_count}/{len(results)} 个文档到 Edit_Doc",
        "results": results,
    })


# ── 批量替换 ─────────────────────────────────────────────────

# @app.route("/api/edit-doc-list")
# def edit_doc_list():
#     """列出 Edit_Doc 中所有可供替换的文档（仅 .docx，python-docx 不支持 .doc）"""
#     files = []
#     if EDIT_DOC_DIR.exists():
#         for d in sorted(EDIT_DOC_DIR.iterdir(), key=lambda x: int(x.name) if x.name.isdigit() else 999):
#             if d.is_dir():
#                 for f in d.iterdir():
#                     if f.suffix.lower() == ".docx":
#                         material_name = ""
#                         for name, num in MATERIAL_DOC_MAP.items():
#                             if num == int(d.name):
#                                 material_name = name
#                                 break
#                         files.append({
#                             "dir_num": int(d.name),
#                             "filename": f.name,
#                             "material": material_name,
#                             "path": f"{d.name}/{f.name}",
#                         })
#     return jsonify(files)


# @app.route("/api/batch-replace", methods=["POST"])
# def batch_replace():
#     """批量替换 Edit_Doc 中文档的文本，保留格式"""
#     data = request.get_json(force=True)
#     rel_path = data.get("file_path", "")
#     replacements = data.get("replacements", {})
#
#     if not rel_path or not replacements:
#         return jsonify({"error": 1, "message": "缺少参数"}), 400
#
#     filepath = EDIT_DOC_DIR / rel_path
#     if not filepath.exists():
#         return jsonify({"error": 1, "message": f"文件不存在: {rel_path}"}), 404
#
#     if filepath.suffix.lower() != ".docx":
#         return jsonify({"error": 1, "message": "仅支持 .docx 格式，.doc 请先用 Word 另存为 .docx"}), 400
#
#     try:
#         doc = Document(str(filepath))
#
#         def replace_in_runs(runs, reps):
#             for run in runs:
#                 for old_text, new_text in reps.items():
#                     if old_text in run.text:
#                         run.text = run.text.replace(old_text, new_text)
#
#         # 段落
#         for paragraph in doc.paragraphs:
#             replace_in_runs(paragraph.runs, replacements)
#
#         # 表格
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for paragraph in cell.paragraphs:
#                         replace_in_runs(paragraph.runs, replacements)
#
#         # 页眉页脚
#         for section in doc.sections:
#             for paragraph in section.header.paragraphs:
#                 replace_in_runs(paragraph.runs, replacements)
#             for paragraph in section.footer.paragraphs:
#                 replace_in_runs(paragraph.runs, replacements)
#
#         # 保存回原文件
#         doc.save(str(filepath))
#
#         count = len(replacements)
#         print(f"[batch-replace] 已替换 {count} 组词 → {filepath}")
#         return jsonify({"error": 0, "message": f"成功替换 {count} 组文本", "count": count})
#
#     except Exception as e:
#         print(f"[batch-replace] 失败: {e}")
#         return jsonify({"error": 1, "message": str(e)}), 500


# ── 参数组管理 ──────────────────────────────────────────────

PARAM_GROUPS_FILE = BASE_DIR / "param_groups.json"


def load_param_groups():
    if PARAM_GROUPS_FILE.exists():
        return json.loads(PARAM_GROUPS_FILE.read_text(encoding="utf-8"))
    # 默认 A/B/C，每个组默认选中所有文档
    defaults = {
        "groups": ["A", "B", "C"],
        "doc_selections": {g: [] for g in ["A", "B", "C"]},
    }
    save_param_groups(defaults)
    return defaults


def save_param_groups(data):
    PARAM_GROUPS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/api/param-groups", methods=["GET"])
def get_param_groups():
    return jsonify(load_param_groups())


@app.route("/api/param-groups", methods=["POST"])
def update_param_groups():
    data = request.get_json(force=True)
    save_param_groups(data)
    return jsonify({"error": 0})

@app.route("/api/param-groups/<group>/folder", methods=["DELETE"])
def delete_group_folder(group):
    """删除模版组的文件夹"""
    import shutil
    group_dir = USER_TEMPLATES_DIR / group
    if group_dir.exists():
        shutil.rmtree(str(group_dir))
    return jsonify({"error": 0})
USER_TEMPLATES_DIR=BASE_DIR/"user_templates"
USER_TEMPLATES_DIR.mkdir(exist_ok=True)
# def get_all_templates(group=None):
#     """所有模板来自 user_templates/ 目录，dir_num 用文件名哈希保证稳定"""
#     result = {}
#     search_dirs = [USER_TEMPLATES_DIR / group] if group else [d for d in USER_TEMPLATES_DIR.glob("*") if d.is_dir()]
#     for d in search_dirs:
#         if d.exists():
#             for f in sorted(d.glob("*.docx")):
#                 result[f.name] = abs(hash(f.name)) % 10000 + 1
#     return result

# from database import get_db, init_db
from pathlib import Path
import hashlib

USER_TEMPLATES_DIR = Path("./user_templates")


def get_stable_hash(filename):
    """使用 MD5 生成稳定哈希（重启不变）"""
    return int(hashlib.md5(filename.encode()).hexdigest()[:8], 16) % 10000 + 1


def get_all_templates(group=None):
    """
    获取所有模板，返回 {文件名: 文件夹编号}
    如果数据库有记录，直接使用；否则创建新映射
    """
    result = {}

    # 1. 确定要扫描的目录
    if group:
        search_dirs = [USER_TEMPLATES_DIR / group]
    else:
        search_dirs = [d for d in USER_TEMPLATES_DIR.glob("*") if d.is_dir()]

    conn = get_db()
    cursor = conn.cursor()

    for d in search_dirs:
        if not d.exists():
            continue

        group_name = d.name  # 目录名作为 group_name

        for f in sorted(d.glob("*.docx")):
            filename = f.name

            # 2. 查询数据库是否已有映射
            cursor.execute('''
                SELECT dir_num FROM template_mappings
                WHERE group_name = ? AND filename = ?
            ''', (group_name, filename))

            row = cursor.fetchone()

            if row:
                # 已有映射，直接使用
                dir_num = row['dir_num']
            else:
                # 新文件，生成稳定哈希作为文件夹编号
                dir_num = get_stable_hash(filename)

                # 保存到数据库
                cursor.execute('''
                    INSERT OR IGNORE INTO template_mappings (group_name, filename, dir_num)
                    VALUES (?, ?, ?)
                ''', (group_name, filename, dir_num))
                conn.commit()

            result[filename] = dir_num

    conn.close()
    return result







#  ''' enumerate()	Python 内置函数，为可迭代对象添加计数器
# 返回	生成 (索引, 元素) 元组的迭代器
# #.glob("*.docx")	返回一个生成器，匹配该目录下所有 .docx 文件
# '''

@app.route("/api/templates",methods=["GET"])
def list_templates():
    """列出所有模板文件名"""
    group=request.args.get('group')
    templates=get_all_templates(group)
    return jsonify(list(templates.keys()))

#在前端用_availableDocs这个列表参数接收。
@app.route("/api/templates/upload",methods=["POST"])
def upload_template():
    group=request.form.get("group","A")
    file=request.files.get("file")
    if not file or not file.filename.endswith(".docx"):
        return jsonify({"error": 1, "message": "请上传 .docx 文件"}), 400
    upload_dir=USER_TEMPLATES_DIR/group
    upload_dir.mkdir(parents=True,exist_ok=True)
    file.save(upload_dir/file.filename)
    return jsonify({"error":0,"filename":file.filename})
@app.route("/api/templates/<path:filename>",methods=["DELETE"])
def delete_template(filename):
    group=request.args.get("group","")
    # 在所有组目录下查找并删除
    for d in USER_TEMPLATES_DIR.glob("*"):
        if d.is_dir():
            fp = d / group/filename
            print(f"删除文件: {fp}")
            if fp.exists():
                fp.unlink()
                return jsonify({"error": 0})
    return jsonify({"error": 0, "message": "文件不存在"})
from flask import Flask, render_template, send_file, request, jsonify, abort
from pathlib import Path
import sys
import tempfile
import subprocess
import os
import shutil


# ========== 打印文档 ==========

@app.route("/print-preview/<int:case_id>/<material_name>")
def print_preview(case_id, material_name):
    """将 docx 转为 HTML 打印预览页"""
    from docx import Document as DocxDocument

    # 直接搜索 Edit_Doc/{case_id}/ 下所有子目录，找匹配的文件
    case_root = EDIT_DOC_DIR / str(case_id)
    if not case_root.exists():
        return "该案件无文档目录", 404

    filepath = None
    for d in sorted(case_root.iterdir(), key=lambda x: x.name):
        if d.is_dir():
            for f in d.glob("*"):
                if get_material_display_name(f.name) == material_name or f.stem == material_name:
                    filepath = f
                    break
            if filepath:
                break

    if not filepath:
        return f"未找到「{material_name}」的文档文件", 404
    try:
        doc = DocxDocument(str(filepath))
    except Exception:
        return "无法打开文档（可能不是 .docx 格式）", 400

    html_parts = []
    for p in doc.paragraphs:
        text = p.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
        if p.style.name.startswith("Heading"):
            level = p.style.name.replace("Heading ", "").replace("Heading", "1")
            html_parts.append(f'<h{level} style="margin:12px 0 4px 0;">{text}</h{level}>')
        elif text.strip():
            html_parts.append(f'<p style="margin:4px 0;text-indent:2em;">{text}</p>')
        else:
            html_parts.append('<br>')

    # 表格
    table_html = []
    for table in doc.tables:
        rows_html = []
        for row in table.rows:
            cells = ''.join(f'<td style="border:1px solid #999;padding:4px 8px;">{cell.text}</td>' for cell in row.cells)
            rows_html.append(f'<tr>{cells}</tr>')
        table_html.append(f'<table style="border-collapse:collapse;width:100%;margin:8px 0;">{"".join(rows_html)}</table>')

    return f'''<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>打印 - {material_name}</title>
<style>@media print{{button{{display:none}}}}body{{font-family:SimSun,serif;max-width:700px;margin:20px auto;padding:0 20px;}}h1{{text-align:center;}}</style>
</head><body>
<h1>{material_name}</h1>
<button onclick="window.print()" style="position:fixed;top:10px;right:10px;padding:10px 20px;background:#8B1A1A;color:#fff;border:none;border-radius:6px;cursor:pointer;font-size:16px;">🖨️ 点击打印</button>
{"".join(html_parts)}
{"".join(table_html)}
<script>setTimeout(function(){{window.print();}},500);</script>
</body></html>'''



def get_material_display_name(fn):
    """模板文件名 → 显示名（去掉 .docx 后缀）"""
    return fn.replace(".docx", "")


@app.route("/api/group-docs/<group>")
def get_group_docs(group):
    """获取某个模版组选中的文档对应的材料名列表"""
    all_t = get_all_templates(group)
    config = load_param_groups()
    docs = config.get("doc_selections", {}).get(group, list(all_t.keys()))
    resolved = [fn for fn in docs if fn in all_t]
    if not resolved:
        resolved = docs  # fallback
    result = [get_material_display_name(fn) for fn in resolved]
    return jsonify(result)

# ── 案件 CRUD API ────────────────────────────────────────────

@app.route("/api/cases", methods=["GET"])
def list_cases():
    db = get_db()
    rows = db.execute("SELECT * FROM cases ORDER BY id DESC").fetchall()
    cases = []
    for row in rows:
        c = dict(row)
        # 查占位符
        ph_rows = db.execute(
            "SELECT param_set, code, value FROM placeholders WHERE case_id=? ORDER BY param_set, code",
            (c["id"],),
        ).fetchall()
        placeholders = {}
        for ph in ph_rows:
            placeholders.setdefault(ph["param_set"], {})[ph["code"]] = ph["value"]
        c["placeholders"] = placeholders
        cases.append(c)
    return jsonify(cases)


@app.route("/api/cases", methods=["POST"])
def create_case():
    data = request.get_json(force=True)
    db = get_db()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur = db.execute(
        "INSERT INTO cases (case_name, subject_name, unit_name, archivist, archive_date, status, reviewer, created_at, active_group) VALUES (?,?,?,?,?,?,?,?,?)",
        (
            data.get("caseName", ""),
            data.get("subjectName", ""),
            data.get("unitName", ""),
            data.get("archivist", ""),
            data.get("archiveDate", now[:10]),
            data.get("status", "active"),
            data.get("reviewer", ""),
            now,
            data.get("active_group", ""),

        ),
    )
    case_id = cur.lastrowid

    # 保存占位符（所有参数组）
    placeholders = data.get("placeholders", {})
    for param_set, vals in placeholders.items():
        for code, value in vals.items():
            if value:
                db.execute(
                    "INSERT INTO placeholders (case_id, param_set, code, value) VALUES (?,?,?,?)",
                    (case_id, param_set, code, value),
                )

    db.commit()
    return jsonify({"error": 0, "id": case_id})


@app.route("/api/cases/<int:case_id>", methods=["GET"])
def get_case(case_id):
    db = get_db()
    row = db.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone()
    if not row:
        abort(404)
    c = dict(row)
    ph_rows = db.execute(
        "SELECT param_set, code, value FROM placeholders WHERE case_id=? ORDER BY param_set, code",
        (case_id,),
    ).fetchall()
    placeholders = {}
    for ph in ph_rows:
        placeholders.setdefault(ph["param_set"], {})[ph["code"]] = ph["value"]
    c["placeholders"] = placeholders
    return jsonify(c)


@app.route("/api/cases/<int:case_id>", methods=["PUT"])
def update_case(case_id):
    data = request.get_json(force=True)
    db = get_db()
    db.execute(
        "UPDATE cases SET case_name=?, subject_name=?, unit_name=?, archivist=?, status=?, reviewer=? ,active_group=? WHERE id=?",
        (
            data.get("caseName", ""),
            data.get("subjectName", ""),
            data.get("unitName", ""),
            data.get("archivist", ""),
            data.get("status", "active"),
            data.get("reviewer", ""),
            data.get("active_group", ""),
            case_id,
        ),
    )

    # 删除旧占位符，重新写入
    db.execute("DELETE FROM placeholders WHERE case_id=?", (case_id,))
    placeholders = data.get("placeholders", {})
    for param_set, vals in placeholders.items():
        for code, value in vals.items():
            if value:
                db.execute(
                    "INSERT INTO placeholders (case_id, param_set, code, value) VALUES (?,?,?,?)",
                    (case_id, param_set, code, value),
                )

    db.commit()
    return jsonify({"error": 0})


@app.route("/api/cases/<int:case_id>", methods=["DELETE"])
def delete_case(case_id):
    db = get_db()
    db.execute("DELETE FROM cases WHERE id=?", (case_id,))
    import shutil
    case_dir=EDIT_DOC_DIR/str(case_id)
    if case_dir.exists():
        shutil.rmtree((case_dir))
    db.commit()
    return jsonify({"error": 0})


# ── 启动 ──────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 5050
    print("=" * 60)
    print("  智慧办案助手 - 离线版")
    print(f"  访问: http://localhost:{port}")
    print("=" * 60)
    app.run(host="0.0.0.0", port=port, debug=True)
