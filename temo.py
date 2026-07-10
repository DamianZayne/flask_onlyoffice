from docx import Document


def batch_replace_keep_format(file_path, replacements, output_path=None):
    """
    批量替换多组文本，完全保留格式

    Args:
        file_path: Word文件路径
        replacements: 字典，{'旧文本1': '新文本1', '旧文本2': '新文本2'}
        output_path: 输出文件路径
    """
    doc = Document(file_path)

    def replace_in_runs(runs, replacements):
        """在runs中执行替换"""
        for run in runs:
            for search_text, replace_text in replacements.items():
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)

    # 处理所有段落
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph.runs, replacements)

    # 处理所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph.runs, replacements)

    # 处理页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_runs(paragraph.runs, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_runs(paragraph.runs, replacements)

    # 保存
    if output_path is None:
        output_path = file_path
    doc.save(output_path)
    print(f"✅ 批量替换完成！")


# 使用示例：同时替换多个词
batch_replace_keep_format(
    file_path="uploads/89c58fca48f8.docx",
    replacements={
        '时间': 'time',
        '监察室': '工商司',
        '议题': '2026年'
    },
    output_path="example_new.docx"
)